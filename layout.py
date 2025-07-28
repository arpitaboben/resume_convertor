import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_colored_underline(paragraph, color=RGBColor(0, 102, 204)):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2]))
    border.append(bottom)
    pPr.append(border)


def add_section_header(doc, text):

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = docx.shared.Inches(0.3)
    table.columns[1].width = docx.shared.Inches(2.5)
    table.columns[2].width = docx.shared.Inches(5.0)

    from docx.enum.table import WD_ROW_HEIGHT_RULE
    table.rows[0].height = docx.shared.Inches(0.18)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Colors
    dark_blue = '20446A'  # RGB(32, 68, 106)
    light_gray = 'DDDDDD'

    # Left cell: dark blue
    cell0 = table.cell(0, 0)
    cell0.text = ''
    cell0._tc.get_or_add_tcPr().append(docx.oxml.parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{dark_blue}"/>'))

    # Middle cell: light gray, section title, bold, spaced
    cell1 = table.cell(0, 1)
    cell1.text = ''
    cell1._tc.get_or_add_tcPr().append(docx.oxml.parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{light_gray}"/>'))
    para = cell1.paragraphs[0]
    run = para.add_run(' '.join(list(text.upper())))
    run.bold = True
    run.font.size = Pt(8)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Remove extra spacing
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1

    # Right cell: dark blue, no text
    cell2 = table.cell(0, 2)
    cell2.text = ''
    cell2._tc.get_or_add_tcPr().append(docx.oxml.parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{dark_blue}"/>'))

    # Remove borders for a cleaner look
    tbl = table._tbl
    for border_dir in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tblPr = tbl.tblPr
        borders = tblPr.xpath('./w:tblBorders')
        if not borders:
            borders = docx.oxml.parse_xml('<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tblPr.append(borders)
        else:
            borders = borders[0]
        border = borders.find(f'w:{border_dir}', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if border is None:
            border = docx.oxml.parse_xml(f'<w:{border_dir} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="nil"/>')
            borders.append(border)
        else:
            border.set(qn('w:val'), 'nil')
    return table


def add_bullet_with_bold_intro(doc, text):
    # Split at first ":" or "." or dash or bolded phrase (heuristic)
    import re
    match = re.match(r"([^.:-]+[:.-]?)\s*(.*)", text)
    if match:
        intro, rest = match.group(1), match.group(2)
    else:
        intro, rest = text, ''
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(intro.strip())
    run.bold = True
    if rest:
        para.add_run(' ' + rest.strip())


def generate_docx(resume_data, output_path="formatted_resume.docx"):
    doc = Document()

    if "name" in resume_data and resume_data["name"]:
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(resume_data["name"])
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph()

   
    if "summary" in resume_data and resume_data["summary"]:
        add_section_header(doc, "Professional Summary")
        summary_text = resume_data["summary"]
        keywords = ["Leadership", "Communication", "Problem Solving", "Analytical", "Teamwork","Project Management", "Creativity", "Adaptability", "Collaboration","Experience", "Achievements", "Accomplished", "Delivered", "Improved","Increased", "Reduced", "Managed", "Developed", "Designed", "Led","Built", "Created", "Implemented", "Organized", "Achieved", "Enhanced","Optimized", "Supported", "Trained", "Innovative", "Strategic","Results-Oriented", "Motivated", "Professional", "Expertise","Certified", "Solution", "Performance", "Growth"]

        para = doc.add_paragraph()
        last_idx = 0
        import re
        for match in re.finditer(r"|".join(map(re.escape, keywords)), summary_text, re.I):
            start, end = match.start(), match.end()
            if start > last_idx:
                para.add_run(summary_text[last_idx:start])
            run = para.add_run(summary_text[start:end])
            run.bold = True
            last_idx = end
        if last_idx < len(summary_text):
            para.add_run(summary_text[last_idx:])
        if "summary_bullets" in resume_data and resume_data["summary_bullets"]:
            for bullet in resume_data["summary_bullets"]:
                add_bullet_with_bold_intro(doc, bullet)

    if "skills" in resume_data and resume_data["skills"]:
        add_section_header(doc, "Technical Skill Sets")
        for skill in sorted(set(resume_data["skills"])):
            doc.add_paragraph(skill, style="List Bullet")

    if ("experience" in resume_data and resume_data["experience"]) or ("projects" in resume_data and resume_data["projects"]):
        add_section_header(doc, "Experience")

        if "experience" in resume_data and resume_data["experience"]:
            for exp in resume_data["experience"]:
                if isinstance(exp, dict):
                    hdr = f"{exp.get('title', '')} ({exp.get('start', '')} - {exp.get('end', '')})"
                    company_loc = f"{exp.get('company', '')}, {exp.get('location', '')}".strip(", ")
                    if hdr.strip():
                        doc.add_paragraph(hdr, style="List Bullet")
                    if company_loc.strip():
                        doc.add_paragraph(company_loc)
                    for detail in exp.get("details", []):
                        if detail.strip():
                            doc.add_paragraph(detail, style="List Bullet 2")
                else:
                    doc.add_paragraph(exp, style="List Bullet")

        if "projects" in resume_data and resume_data["projects"]:
            for proj in resume_data["projects"]:
                if isinstance(proj, dict):
                    line = f"{proj.get('name', '')}: {proj.get('description', '')}".strip(": ")
                else:
                    line = proj
                if line and len(line.strip()) > 3 and not line.strip().startswith("•"):
                    doc.add_paragraph(line, style="List Bullet")


    if "certifications" in resume_data and resume_data["certifications"]:
        add_section_header(doc, "Certifications")
        for cert in resume_data["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    if "achievements" in resume_data and resume_data["achievements"]:
        add_section_header(doc, "Achievements")
        for ach in resume_data["achievements"]:
            doc.add_paragraph(ach, style="List Bullet")


    if "education" in resume_data and resume_data["education"]:
        add_section_header(doc, "Academic Details")
        for edu in resume_data["education"]:
            if isinstance(edu, dict):
                parts = [edu.get("degree", ""), edu.get("institution", ""), edu.get("major", ""), edu.get("score", "")]
                line = ", ".join([part for part in parts if part])
            else:
                line = edu
            if line:
                doc.add_paragraph(line, style="List Bullet")

    doc.save(output_path)
    return output_path
