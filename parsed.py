import os
import re
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from pyresparser import ResumeParser as PyResParser

POPPLER_PATH = r"C:\poppler-24.08.0\poppler-24.08.0\Library\bin"
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

def save_uploaded_file(uploaded_file, save_dir="temp_resumes"):
    os.makedirs(save_dir, exist_ok=True)
    file_path = os.path.join(save_dir, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path

def clean_summary(summary):
    summary = re.sub(r'Address:.*?(?=\\.|$)', '', summary, flags=re.I)
    summary = re.sub(r'[~=:;*><,\.]{2,}', ' ', summary)
    summary = re.sub(r'\bSe CONTACT\b', '', summary, flags=re.I)
    summary = re.sub(r'=f Zi > i', '', summary)
    summary = re.sub(r'\d+\s*innovative software products\.', 'innovative software products.', summary)
    summary = re.sub(r'\s+', ' ', summary)
    return summary.strip()

def extract_email(text):
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.\w+", text)
    return match.group(0) if match else ""

def extract_name(text, docx_path=None):
    match = re.search(r'(?i)name[:\s-]+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)+)', text)
    if match:
        return match.group(1).strip()
    lines = text.splitlines()
    for line in lines:
        line = line.strip()
        if line:
            return line

    if docx_path:
        try:
            from docx import Document
            doc = Document(docx_path)
            for section in doc.sections:
                header = section.header
                for para in header.paragraphs:
                    header_line = para.text.strip()
                    if header_line:
                        return header_line
        except Exception as e:
            print("Error extracting header text:", e)
    return ""

def extract_skills(text):
    keywords = [
        "JavaScript", "React", "Node", "Express", "MongoDB", "SQL", "HTML", "CSS", "AWS", "Redux", "PHP", "JSP",
        "Bootstrap", "JQuery", "MySQL", "Python", "R", "Shiny", "TypeScript", "Django", "Flask", "Agile", "Scrum",
        "Tensorflow", "Machine learning", "Opencv", "Api", "Sqlalchemy", "Tableau", "Github", "Modeling", "Analysis",
        "Ai", "Data analysis", "Health", "System", "Ibm", "Startup", "Prototyping", "Certification", "Strategy"
    ]
    skills = set()
    for kw in keywords:
        if re.search(r'\b' + re.escape(kw) + r'\b', text, re.I):
            skills.add(kw)
    return sorted(skills)

def extract_contact(text):
    contact = []
    phone = re.search(r"(\+91[-\s]?\d{10}|\b\d{10}\b)", text)
    if phone:
        contact.append(f"Phone: {phone.group(0)}")
    addr = re.search(r"Address: (.+)", text)
    if addr:
        contact.append(addr.group(0))
    linkedin = re.search(r"(https?://(www\.)?linkedin\.com/[^\s]+)", text)
    if linkedin:
        contact.append(f"LinkedIn: {linkedin.group(0)}")
    return contact

def parse_structured_resume(text, existing_json=None):
    clean_text = re.sub(r"[^\x00-\x7F]+", " ", text)
    lines = [line.strip("•¢©- ") for line in clean_text.splitlines() if line.strip()]

    # Enhanced section names and their triggers with more variations
    section_names = [
        ("summary", ["summary", "objective", "profile", "about"]),
        ("education", ["education", "academic", "academic details", "qualification", "degree"]),
        ("certifications", ["certification", "certifications", "certificates", "certified"]),
        ("projects", ["projects", "project", "portfolio", "achievements", "key achievements", "projects and certifications"]),
        ("experience", ["experience", "work experience", "professional experience", "internship", "employment", "work history"]),
        ("skills", ["skills", "technical skills", "competencies", "technologies"]),
    ]
    section_map = {}
    current_section = None
    buffer = []

    for line in lines:
        if '|' in line:
            left, right = [part.strip() for part in line.split('|', 1)]
            l_line = left.lower()
            matched = False
            for key, triggers in section_names:
                if any(l_line.startswith(t) for t in triggers):
                    if current_section and buffer:
                        section_map[current_section] = buffer
                    current_section = key
                    buffer = [right]
                    matched = True
                    break
            if matched:
                continue 
        l_line = line.lower().strip()
        matched = False
        
        for key, triggers in section_names:
            if any(l_line.startswith(t) for t in triggers):
                if current_section and buffer:
                    section_map[current_section] = buffer
                current_section = key
                buffer = []
                matched = True
                break
          
            elif any(f"{t}:" in l_line or f"{t} -" in l_line or f"{t}." in l_line for t in triggers):
                if current_section and buffer:
                    section_map[current_section] = buffer
                current_section = key
                buffer = []
                matched = True
                break
            
            elif any(t.upper() in line.upper() and len(line.strip()) < 20 for t in triggers):
                if current_section and buffer:
                    section_map[current_section] = buffer
                current_section = key
                buffer = []
                matched = True
                break
        if not matched:
            buffer.append(line)
    if current_section and buffer:
        section_map[current_section] = buffer

  
    exp_lines = section_map.get("experience", [])
    experiences = []
    exp = {}
    job_title_pattern = re.compile(r"^(?P<title>.+?),\s*(?P<start>\d{2}/\d{4})\s*-\s*(?P<end>Current|\d{2}/\d{4})")
    company_pattern = re.compile(r"^(?P<company>.+?)(?:\s*-\s*(?P<location>.+))?$")
    i = 0
    while i < len(exp_lines):
        line = exp_lines[i]
        m = job_title_pattern.match(line)
        if m:
            if exp:
                experiences.append(exp)
            exp = {
                "title": m.group("title").strip(),
                "start": m.group("start"),
                "end": m.group("end") if m.group("end") != "Current" else None,
                "company": "",
                "location": "",
                "details": []
            }
            if i + 1 < len(exp_lines):
                next_line = exp_lines[i + 1].strip()
                m2 = company_pattern.match(next_line)
                if m2:
                    exp["company"] = m2.group("company").strip()
                    exp["location"] = m2.group("location").strip() if m2.group("location") else ""
                    i += 1  
            i += 1
   
            while i < len(exp_lines) and not job_title_pattern.match(exp_lines[i]):
                if exp_lines[i].strip():
                    exp.setdefault("details", []).append(exp_lines[i].strip())
                i += 1
        else:
            i += 1
    if exp:
        experiences.append(exp)

    summary = clean_summary(" ".join(section_map.get("summary", [])))

    summary = re.split(r'contact', summary, flags=re.I)[0].strip()

    summary = re.sub(r'([A-Za-z])(?:\s*\1){2,}$', '', summary).strip()
    raw_lines = summary.splitlines()
    cleaned_lines = []
    for line in raw_lines:
        l = line.strip()
        if not l:
            continue
        if "contact" in l.lower():
            continue
        if re.match(r"^(.)\1{4,}$", l):  
            continue
        if not re.search(r"[a-zA-Z]", l): 
            continue
        cleaned_lines.append(l)
    summary = "\n".join(cleaned_lines)


    education_lines = section_map.get("education", [])
    educations = []
    edu_entry = {}
    project_content = []
    experience_content = []
    
    for line in education_lines:

        if any(edu_term in line for edu_term in ["M.Tech", "B.E", "B.Tech", "Bachelor", "Master"]):
            if edu_entry and ("degree" in edu_entry or "institution" in edu_entry):
                educations.append(edu_entry)
            edu_entry = {"degree": line}
        elif any(edu_term in line for edu_term in ["University", "Institute", "School", "College"]):
            edu_entry["institution"] = line
        elif any(score_term in line for score_term in ["CGPA", "%", "Grade"]):
            edu_entry["score"] = line
        elif "Major" in line or "major" in line:
            edu_entry["major"] = line
        
        elif any(project_term in line for project_term in ["Hackathon", "Project", "Developed", "Built", "Created", "YOLO", "ML algorithms", "API", "Flask", "Backend"]):
            project_content.append(line)
    
        elif any(exp_term in line for exp_term in ["Led", "Team", "Member", "Core Member", "Leadership", "Community"]):
            experience_content.append(line)
        else:
            if "degree" in edu_entry or "institution" in edu_entry:
                edu_entry.setdefault("details", []).append(line)
    
    if edu_entry:
        educations.append(edu_entry)

    certification_lines = section_map.get("certifications", [])
    certifications = []
    for line in certification_lines:
        if line:
            certifications.append(line)


    project_lines = section_map.get("projects", []) + project_content 
    projects = []
    project = {}
    
    for line in project_lines:
        m = re.match(r"^(.*?):\s+(.*)", line)
        if m:
            if project:
                projects.append(project)
            project = {"name": m.group(1).strip("- "), "description": m.group(2)}
        elif line.startswith("-") or line.startswith("•"):
            if project:
                projects.append(project)
            project = {"name": line.lstrip("-• ").strip(), "description": ""}

        elif any(project_term in line for project_term in ["Hackathon", "Developed", "Built", "Created", "YOLO", "ML algorithms"]):
            if project:
                projects.append(project)
            project = {"name": line.strip(), "description": ""}
        else:
            if project:
                if project.get("description"):
                    project["description"] += " " + line
                else:
                    project["description"] = line
    if project:
        projects.append(project)

    email = extract_email(text)
    address = None
    adr_match = re.search(r"Address: (.+)", text)
    if adr_match:
        address = adr_match.group(1).strip()

    skills = extract_skills(text)
    name = extract_name(text, docx_path="path/to/resume.docx")
    contact = extract_contact(text)

    result = {}
    if existing_json:
        result.update(existing_json)
    result["name"] = name or ""
    result["summary"] = summary or ""
    result["address"] = address or ""
    result["email"] = email or ""
    result["skills"] = skills if skills is not None else []
    result["contact"] = contact if contact is not None else []
    result["education"] = educations if educations else []
    result["certifications"] = certifications if certifications else []
    result["projects"] = projects if projects else []
    result["experience"] = experiences if experiences else []
    for field in ["experience", "education", "projects", "skills", "contact", "certifications"]:
        if result.get(field) is None:
            result[field] = []
    return result

def extract_text_from_docx(docx_path):
    from docx import Document
    doc = Document(docx_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)


def parse_resume(file_path):
    import os
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        text = extract_text_from_docx(file_path)
        if text and len(text.strip()) > 20:
            parsed = parse_structured_resume(text)
            for field in ["experience", "education", "projects", "skills", "contact", "certifications"]:
                if parsed.get(field) is None:
                    parsed[field] = []
            return parsed
    print("Tesseract version:", pytesseract.get_tesseract_version())
    text = ""
    # 1. Try direct PDF text extraction
    try:
        from PyPDF2 import PdfReader
        print("📄 Trying direct PDF text extraction with PyPDF2...")
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""
        if text and len(text.strip()) > 100:  # Heuristic: enough text
            print("✅ Extracted text from PDF directly.")
            parsed = parse_structured_resume(text)
            for field in ["experience", "education", "projects", "skills", "contact", "certifications"]:
                if parsed.get(field) is None:
                    parsed[field] = []
            return parsed
        else:
            print("⚠️ Direct PDF text extraction yielded little or no text, trying OCR fallback.")
    except Exception as e:
        print("⚠️ Direct PDF text extraction failed:", e)
    
    # 2. OCR fallback
    print("🧠 Running OCR fallback...")
    try:
        images = convert_from_path(file_path, poppler_path=POPPLER_PATH, dpi=300)
        # Save the first page image for debugging
        if images:
            images[0].save("first_page.png")
        text = ""
        for img in images:
            img = img.convert('L')  # Grayscale for better OCR
            ocr_text = pytesseract.image_to_string(img, config='--psm 3')
            print("FULL OCR TEXT:\n", ocr_text)
            # Try to extract the name
            for line in ocr_text.splitlines():
                words = line.split()
                if len(words) >= 2 and sum(w[0].isupper() for w in words if w and w[0].isalpha()) >= 2:
                    print("Likely name:", line)
                    break
            text += ocr_text
        parsed = parse_structured_resume(text)
        for field in ["experience", "education", "projects", "skills", "contact", "certifications"]:
            if parsed.get(field) is None:
                parsed[field] = []
        return parsed
    except Exception as e:
        print(" OCR fallback failed:", e)
    
    # 3. Last resort: PyResParser
    try:
        print("Trying PyResParser as last resort...")
        data = PyResParser(file_path).get_extracted_data()
        if data:
            for field in ["experience", "education", "projects", "skills", "contact", "certifications"]:
                if data.get(field) is None:
                    data[field] = []
            return data
    except Exception as e:
        print(" PyResParser failed:", e)
    return {"error": "All extraction methods failed."}

if __name__ == "__main__":
    import json
    with open("resume.txt", "r", encoding="utf-8") as f:
        text = f.read()
    parsed = parse_structured_resume(text)
    with open("resume_parsed.json", "w", encoding="utf-8") as f:
        json.dump(parsed, f, indent=2)
    print(json.dumps(parsed, indent=2))
